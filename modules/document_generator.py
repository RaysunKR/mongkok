import os
import markdown
from datetime import datetime
from typing import Dict, Any, Optional, List
from jinja2 import Template
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from ..utils.logger import get_logger
from ..utils.config_loader import ConfigLoader


class DocumentGenerator:
    """文档生成模块 - 支持多种格式文档生成"""

    def __init__(self, config: ConfigLoader = None):
        self._config = config or ConfigLoader()
        self._logger = get_logger('DocumentGenerator')
        self._output_dir = os.path.join(
            self._config.storage.get('data_dir', 'data'),
            'documents'
        )
        os.makedirs(self._output_dir, exist_ok=True)

        capabilities = self._config.capabilities
        self._enabled = capabilities.get('document_generation', True)

    async def generate_markdown(
        self,
        content: str,
        title: str = "文档",
        metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """生成Markdown文档"""
        if not self._enabled:
            return {'success': False, 'error': '文档生成功能未启用'}

        try:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"{title}_{timestamp}.md"
            filepath = os.path.join(self._output_dir, filename)

            # 构建Markdown内容
            metadata = metadata or {}
            markdown_content = f"""# {title}

**生成时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

---

{content}

---

## 文档元数据
"""

            for key, value in metadata.items():
                markdown_content += f"- **{key}**: {value}\n"

            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(markdown_content)

            self._logger.info(f"Markdown文档生成成功: {filepath}")

            return {
                'success': True,
                'filepath': filepath,
                'filename': filename,
                'format': 'markdown',
                'size': os.path.getsize(filepath)
            }

        except Exception as e:
            self._logger.error(f"Markdown文档生成失败: {e}")
            return {'success': False, 'error': str(e)}

    async def generate_html(
        self,
        content: str,
        title: str = "文档",
        style: Optional[str] = None
    ) -> Dict[str, Any]:
        """生成HTML文档"""
        if not self._enabled:
            return {'success': False, 'error': '文档生成功能未启用'}

        try:
            # 转换Markdown到HTML
            html_content = markdown.markdown(content, extensions=['tables', 'fenced_code'])

            # 默认样式
            if not style:
                style = """
                <style>
                    body {
                        font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, oxygen, Ubuntu, sans-serif;
                        line-height: 1.6;
                        max-width: 1200px;
                        margin: 0 auto;
                        padding: 20px;
                        color: #333;
                    }
                    h1, h2, h3 {
                        color: #2c3e50;
                        border-bottom: 2px solid #3498db;
                        padding-bottom: 10px;
                    }
                    table {
                        width: 100%;
                        border-collapse: collapse;
                        margin: 20px 0;
                    }
                    table, th, td {
                        border: 1px solid #ddd;
                        padding: 12px;
                    }
                    th {
                        background-color: #3498db;
                        color: white;
                    }
                    code {
                        background-color: #f4f4f4;
                        padding: 2px 6px;
                        border-radius: 3px;
                        font-family: 'Courier New', monospace;
                    }
                    pre {
                        background-color: #2d3436;
                        color: #dfe6e9;
                        padding: 20px;
                        border-radius: 5px;
                        overflow-x: auto;
                    }
                    pre code {
                        background-color: transparent;
                        color: inherit;
                    }
                </style>
                """

            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"{title}_{timestamp}.html"
            filepath = os.path.join(self._output_dir, filename)

            full_html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    {style}
</head>
<body>
    <h1>{title}</h1>
    <p><em>生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</em></p>
    <hr>
    {html_content}
</body>
</html>"""

            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(full_html)

            self._logger.info(f"HTML文档生成成功: {filepath}")

            return {
                'success': True,
                'filepath': filepath,
                'filename': filename,
                'format': 'html',
                'size': os.path.getsize(filepath)
            }

        except Exception as e:
            self._logger.error(f"HTML文档生成失败: {e}")
            return {'success': False, 'error': str(e)}

    async def generate_word_document(
        self,
        content: str,
        title: str = "文档",
        metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """生成Word文档"""
        if not self._enabled:
            return {'success': False, 'error': '文档生成功能未启用'}

        try:
            doc = Document()

            # 添加标题
            title_heading = doc.add_heading(title, 0)
            title_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # 添加元数据
            if metadata:
                metadata_para = doc.add_paragraph()
                metadata_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                metadata_para.add_run(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                for key, value in metadata.items():
                    metadata_para.add_run(f"{key}: {value}\n")
                metadata_para.runs[0].font.size = Pt(10)

            doc.add_paragraph()
            doc.add_paragraph("=" * 50)
            doc.add_paragraph()

            # 添加内容（假设是Markdown格式）
            lines = content.split('\n')
            current_list = None

            for line in lines:
                line = line.strip()
                if not line:
                    doc.add_paragraph()
                    continue

                # 标题
                if line.startswith('#'):
                    level = min(len(line) - len(line.lstrip('#')), 3)
                    heading_text = line.lstrip('#').strip()
                    doc.add_heading(heading_text, level=level)
                # 列表
                elif line.startswith('- ') or line.startswith('* '):
                    if current_list is None:
                        para = doc.add_paragraph(line[2:], style='List Bullet')
                    else:
                        para.add_run(line[2:])
                    current_list = doc.add_paragraph(line[2:], style='List Bullet')
                # 代码块
                elif line.startswith('```'):
                    continue
                # 普通段落
                else:
                    doc.add_paragraph(line)

            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"{title}_{timestamp}.docx"
            filepath = os.path.join(self._output_dir, filename)

            doc.save(filepath)

            self._logger.info(f"Word文档生成成功: {filepath}")

            return {
                'success': True,
                'filepath': filepath,
                'filename': filename,
                'format': 'docx',
                'size': os.path.getsize(filepath)
            }

        except Exception as e:
            self._logger.error(f"Word文档生成失败: {e}")
            return {'success': False, 'error': str(e)}

    async def generate_pdf(
        self,
        content: str,
        title: str = "文档",
        page_size: str = "A4"
    ) -> Dict[str, Any]:
        """生成PDF文档"""
        if not self._enabled:
            return {'success': False, 'error': '文档生成功能未启用'}

        try:
            # 转换Markdown到HTML
            html_content = markdown.markdown(content, extensions=['tables', 'fenced_code'])

            page_size_obj = A4 if page_size.upper() == 'A4' else letter

            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"{title}_{timestamp}.pdf"
            filepath = os.path.join(self._output_dir, filename)

            doc = SimpleDocTemplate(filepath, pagesize=page_size_obj)
            styles = getSampleStyleSheet()

            story = []

            # 标题
            title_style = styles['Heading1']
            title_style.alignment = 1  # 居中
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 12))

            # 日期
            story.append(Paragraph(
                f"<i>生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</i>",
                styles['Normal']
            ))
            story.append(Spacer(1, 12))
            story.append(Spacer(1, 12))

            # 分割线
            story.append(Paragraph("=" * 50, styles['Normal']))
            story.append(Spacer(1, 12))

            # 处理内容
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    story.append(Spacer(1, 6))
                    continue

                if line.startswith('###'):
                    story.append(Paragraph(line[3:].strip(), styles['Heading3']))
                    story.append(Spacer(1, 6))
                elif line.startswith('##'):
                    story.append(Paragraph(line[2:].strip(), styles['Heading2']))
                    story.append(Spacer(1, 6))
                elif line.startswith('#'):
                    story.append(Paragraph(line[1:].strip(), styles['Heading1']))
                    story.append(Spacer(1, 6))
                else:
                    story.append(Paragraph(line, styles['Normal']))
                    story.append(Spacer(1, 6))

            doc.build(story)

            self._logger.info(f"PDF文档生成成功: {filepath}")

            return {
                'success': True,
                'filepath': filepath,
                'filename': filename,
                'format': 'pdf',
                'size': os.path.getsize(filepath)
            }

        except Exception as e:
            self._logger.error(f"PDF文档生成失败: {e}")
            return {'success': False, 'error': str(e)}

    async def generate_from_template(
        self,
        template_path: str,
        data: Dict[str, Any],
        output_format: str = "markdown"
    ) -> Dict[str, Any]:
        """从模板生成文档"""
        if not self._enabled:
            return {'success': False, 'error': '文档生成功能未启用'}

        try:
            with open(template_path, 'r', encoding='utf-8') as f:
                template_content = f.read()

            template = Template(template_content)
            rendered_content = template.render(**data)

            # 将内容转换为HTML用于进一步处理
            html_content = markdown.markdown(rendered_content)

            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            title = data.get('title', '文档')
            filename = f"{title}_{timestamp}.{output_format}"
            filepath = os.path.join(self._output_dir, filename)

            if output_format == 'html':
                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write(html_content)
            else:
                # 其他格式需要进一步处理
                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write(rendered_content)

            self._logger.info(f"模板文档生成成功: {filepath}")

            return {
                'success': True,
                'filepath': filepath,
                'filename': filename,
                'format': output_format
            }

        except Exception as e:
            self._logger.error(f"模板文档生成失败: {e}")
            return {'success': False, 'error': str(e)}

    async def generate_report(
        self,
        data: List[Dict[str, Any]],
        title: str = "报告",
        summary: Optional[str] = None
    ) -> Dict[str, Any]:
        """生成结构化报告"""
        if not self._enabled:
            return {'success': False, 'error': '文档生成功能未启用'}

        try:
            # 构建Markdown报告
            report_content = f"""# {title}

**生成时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

"""

            if summary:
                report_content += f"""## 摘要

{summary}

---

"""

            report_content += "## 详细内容\n\n"

            for i, item in enumerate(data, 1):
                report_content += f"### 项目 {i}\n\n"
                for key, value in item.items():
                    report_content += f"**{key}**: {value}\n\n"
                report_content += "---\n\n"

            # 生成Markdown
            return await self.generate_markdown(
                report_content,
                title,
                {'type': 'report', 'items_count': len(data)}
            )

        except Exception as e:
            self._logger.error(f"报告生成失败: {e}")
            return {'success': False, 'error': str(e)}
